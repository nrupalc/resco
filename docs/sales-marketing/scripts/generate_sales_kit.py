from __future__ import annotations

from pathlib import Path
from shutil import copyfile
import subprocess
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[3]
SALES_ROOT = ROOT / "docs" / "sales-marketing"
OUT = SALES_ROOT / "generated"
WORD_DIR = OUT / "word"
PDF_DIR = OUT / "pdf"
IMG_DIR = OUT / "images"
EXCEL_DIR = OUT / "excel"
RENDER_DIR = OUT / "rendered"

LOGO = ROOT / "docs" / "brand" / "logos" / "final" / "bright-roof-final-logo-transparent.png"
RENDER = Path("/Users/nrupal/.codex/plugins/cache/openai-primary-runtime/documents/26.521.10419/skills/documents/render_docx.py")
PY = Path("/Users/nrupal/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3")

NAVY = "0E1E4D"
AMBER = "F5A524"
TERRACOTTA = "C75B3C"
CREAM = "FAF6EE"
SLATE = "2C3444"
MUTED = "6B7280"
LIGHT = "E8EBEF"
WHITE = "FFFFFF"


CONTACT = {
    "website": "www.brightroofpower.com",
    "email": "support@brightroofpower.com",
    "phone": "+91 93902 10407",
    "address": "8-3-945/8/18&19 Pancom Business Centre, Ameerpet, Hyderabad, 500073, Telangana",
}

PEOPLE = [
    ("Srinivas Koppada", "Technical & Project Coordination", CONTACT["phone"]),
    ("Vijay", "Society Relations", CONTACT["phone"]),
    ("Kalyan", "Society Outreach", CONTACT["phone"]),
]


def ensure_dirs() -> None:
    for d in [WORD_DIR, PDF_DIR, IMG_DIR, EXCEL_DIR, RENDER_DIR]:
        d.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = LIGHT) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(SLATE)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Title", 22, NAVY, 0, 6),
        ("Heading 1", 15, NAVY, 14, 7),
        ("Heading 2", 12.5, NAVY, 10, 5),
        ("Heading 3", 11.5, SLATE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name != "Title"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Bright Roof Power Systems")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run.font.size = Pt(9)
    header.add_run("  |  Solar that stays.").font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"{CONTACT['website']}  |  {CONTACT['email']}  |  {CONTACT['phone']}")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_title_block(doc: Document, title: str, subtitle: str | None = None) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(LOGO), width=Inches(1.05))
    p = doc.add_paragraph()
    p.style = "Title"
    p.add_run(title)
    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(8)
        r = s.add_run(subtitle)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(10)
    rr = rule.add_run(" " * 14)
    rr.font.highlight_color = None
    rr.font.color.rgb = RGBColor.from_string(AMBER)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), AMBER)
    border.append(bottom)
    rule._p.get_or_add_pPr().append(border)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    if widths:
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], LIGHT)
        set_cell_border(hdr[i])
        set_cell_margins(hdr[i])
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(NAVY)
                run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cells[i])
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.25)
                    run.font.color.rgb = RGBColor.from_string(SLATE)
    doc.add_paragraph()
    return table


def save_doc(doc: Document, filename: str) -> Path:
    path = WORD_DIR / filename
    doc.save(path)
    return path


def new_doc(title: str, subtitle: str | None = None) -> Document:
    doc = Document()
    style_doc(doc, title)
    add_title_block(doc, title, subtitle)
    return doc


def leaflet() -> Path:
    doc = new_doc("Rooftop solar for apartment societies", "No upfront installation cost for the society.")
    p = doc.add_paragraph()
    p.add_run("Bright Roof installs, owns, operates, and maintains rooftop solar systems for apartment societies in Hyderabad. ").bold = True
    p.add_run("Your society pays only for solar electricity consumed or offset under the signed agreement, with a starting proposal of 10% saving against the agreed electricity-tariff benchmark.")
    add_table(doc, ["Item", "Bright Roof position"], [
        ["Upfront installation cost", "Nil for the society"],
        ["Ownership", "Bright Roof owns the system"],
        ["Maintenance", "Bright Roof handles operation, cleaning, monitoring, and repairs"],
        ["Monthly payment", "Society pays for solar electricity under the agreement"],
        ["Saving", "Starting proposal of 10% against agreed electricity benchmark"],
        ["Grid connection", "Existing grid supply remains for balance electricity"],
        ["Survey", "Done only after society permission"],
    ], [2.2, 4.0])
    doc.add_heading("Why societies consider this model", level=1)
    for item in [
        "No lakhs of rupees collected from residents for installation.",
        "No long-term maintenance burden on the committee.",
        "No automatic handover of operational risk after payback.",
        "Lower common-area electricity cost without changing normal grid supply.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("First step", level=1)
    for item in [
        "Share the last 12 months of common-area electricity bills.",
        "Confirm consumer service number and sanctioned load details.",
        "Allow rooftop and electrical-room access for a technical survey.",
        "Share society registration and authorized signatory process.",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph(f"Contact: {CONTACT['email']} | {CONTACT['phone']} | {CONTACT['website']}")
    return save_doc(doc, "01-bright-roof-society-leaflet.docx")


def company_profile() -> Path:
    doc = new_doc("Bright Roof Power Systems", "Short company profile for apartment society committees.")
    add_table(doc, ["Field", "Detail"], [
        ["Legal name", "M/s. Bright Roof Power Systems"],
        ["Short name", "Bright Roof"],
        ["Business model", "Rooftop solar PPA / RESCO model"],
        ["Location", "Hyderabad, Telangana"],
        ["Website", CONTACT["website"]],
        ["Email", CONTACT["email"]],
        ["Phone", CONTACT["phone"]],
        ["Address", CONTACT["address"]],
    ], [1.8, 4.5])
    doc.add_heading("What Bright Roof Does", level=1)
    doc.add_paragraph("Bright Roof Power Systems is a Hyderabad-based rooftop solar company focused on apartment societies. We install, own, operate, and maintain rooftop solar systems. The society does not pay the installation cost upfront.")
    doc.add_heading("Who We Serve", level=1)
    for item in ["Lifts", "Pumps", "Common lighting", "Clubhouse", "STP", "Security and maintenance loads", "Other shared facilities"]:
        add_bullet(doc, item)
    doc.add_heading("How The Process Works", level=1)
    add_table(doc, ["Step", "What happens"], [
        ["1. Bill review", "Bright Roof reviews 12 months of electricity bills and tariff details."],
        ["2. Site survey", "Rooftop, electrical room, meter location, and cable route are inspected."],
        ["3. Proposal", "Bright Roof prepares a society-specific commercial and technical proposal."],
        ["4. Agreement", "Final terms are documented in a Power Purchase Agreement."],
        ["5. Installation", "Bright Roof coordinates installation, safety, and approvals."],
        ["6. Operation", "Bright Roof monitors, maintains, cleans, and supports the system."],
    ], [1.5, 4.8])
    return save_doc(doc, "02-bright-roof-company-profile.docx")


def document_checklist() -> Path:
    doc = new_doc("Document Request Checklist", "For first-stage Bright Roof review.")
    sections = {
        "Electricity Details": ["Last 12 months of common-area electricity bills.", "Consumer service number(s).", "Sanctioned load details, if shown separately.", "Current tariff category, if known.", "Any existing net-metering or solar application history."],
        "Society Details": ["Exact registered name of the society.", "Society registration certificate.", "Society bye-laws, if available.", "Current committee office-bearer list.", "Authorized signatory details."],
        "Site Access Details": ["Rooftop access contact person.", "Electrical room access contact person.", "Meter location access contact person.", "Preferred day/time for inspection.", "Any known roof leakage, structural, or access concerns."],
    }
    for heading, items in sections.items():
        doc.add_heading(heading, level=1)
        for item in items:
            add_bullet(doc, "[ ] " + item)
    return save_doc(doc, "03-document-request-checklist.docx")


def consent_form() -> Path:
    doc = new_doc("Permission For Bill Review And Rooftop/Electrical Survey", "Preliminary review consent. Requires final legal review before formal use.")
    doc.add_paragraph("Date: ____________________")
    add_table(doc, ["Field", "Details"], [
        ["Society name", ""],
        ["Society address", ""],
        ["Contact person", ""],
        ["Designation", ""],
        ["Phone", ""],
        ["Email", ""],
    ], [1.8, 4.5])
    doc.add_heading("Permission Granted", level=1)
    for item in [
        "Review electricity bills and related consumption/tariff information shared by the society.",
        "Review society registration, office-bearer, and signatory information shared by the society.",
        "Visit and inspect the rooftop, electrical room, meter area, cable route, and related common areas by prior appointment.",
        "Take site photographs and notes required for technical and commercial feasibility review.",
        "Prepare a preliminary proposal or draft documentation for committee review.",
    ]:
        add_number(doc, item)
    doc.add_heading("Clarifications", level=1)
    for item in [
        "This permission is only for preliminary review and survey.",
        "This permission does not create an obligation for the society to sign a final agreement.",
        "Bright Roof will not begin installation work based only on this consent.",
        "Final terms will be governed by separately approved and signed documents.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Signatures", level=1)
    add_table(doc, ["Society representative", "Bright Roof representative"], [
        ["Name:\nDesignation:\nSignature:\nPhone:\nEmail:", "Name:\nSignature:\nPhone:\nEmail: support@brightroofpower.com"],
    ], [3.1, 3.1])
    return save_doc(doc, "04-bill-sharing-and-survey-consent.docx")


def faq() -> Path:
    doc = new_doc("FAQ And Objection Replies", "Internal field reference for Srinivas, Vijay, and Kalyan.")
    qa = [
        ("Is installation really free?", "There is no upfront installation payment from the society. Bright Roof funds, owns, operates, and maintains the rooftop solar system. The society pays only for solar electricity under the signed agreement."),
        ("What saving do we get?", "The starting proposal is 10% saving against the agreed electricity-tariff benchmark. Bright Roof confirms the exact benchmark and proposal after reviewing bills and site feasibility."),
        ("Can you give more than 10%?", "Possibly, but do not commit before review. Bright Roof can check whether a higher saving is possible after seeing the bills, rooftop feasibility, and expected system size."),
        ("What is the installation cost?", "Bright Roof funds the system. The society's decision is about the tariff saving and agreement terms, not paying the installation cost upfront."),
        ("Who owns the system?", "Bright Roof owns the system during the agreement period."),
        ("Will you hand over the system after payback?", "No automatic handover is part of the standard Bright Roof model. Bright Roof remains responsible for the system, so the society does not inherit operational risk."),
        ("What system size will be installed?", "That depends on electricity consumption, sanctioned load, rooftop space, electrical connection, and survey results."),
        ("What if the roof is damaged?", "The final agreement and installation method will address safe installation, access rules, roof protection, insurance, and responsibility for damage caused by Bright Roof or its installers."),
        ("Are we signing anything today?", "No final PPA is required in the first meeting. The first step is permission for bill review and rooftop/electrical survey."),
    ]
    for q, a in qa:
        doc.add_heading(q, level=2)
        doc.add_paragraph(a)
    doc.add_heading("Escalate before answering in writing", level=1)
    for item in ["Exact tariff commitment.", "Discount above 10%.", "Final system capacity.", "Contract duration changes.", "Handover after payback.", "Roof damage liability wording.", "Payment default consequences.", "Net-metering guarantee."]:
        add_bullet(doc, item)
    return save_doc(doc, "05-faq-and-objection-replies.docx")


def comparison() -> Path:
    doc = new_doc("Bright Roof Vs Alternatives", "Three ways an apartment society can go solar.")
    add_table(doc, ["Option", "Upfront cost", "Ownership", "Maintenance", "Best for"], [
        ["Society buys the system", "Paid by society/residents", "Society", "Society manages", "Societies comfortable collecting capex and managing the system"],
        ["Payback-and-handover model", "Usually low or nil depending on model", "May transfer later", "May shift to society", "Societies that want eventual ownership"],
        ["Bright Roof model", "Nil for society", "Bright Roof", "Bright Roof handles it", "Societies that want savings without upfront collection or maintenance burden"],
    ], [1.45, 1.2, 1.0, 1.25, 1.85])
    doc.add_heading("Simple Decision", level=1)
    doc.add_paragraph("Choose society-owned solar if residents want to fund and manage the system. Choose Bright Roof if the society wants lower electricity cost without upfront installation payment and without taking over long-term system maintenance.")
    return save_doc(doc, "06-bright-roof-vs-alternatives.docx")


def pitch_outline() -> Path:
    doc = new_doc("Committee Pitch Deck", "Six-slide speaking outline.")
    slides = [
        ("1. Bright Roof Power Systems", "Rooftop solar for apartment societies, without upfront installation cost.", ["Bright Roof installs, owns, operates, and maintains rooftop solar systems.", "Starting proposal: 10% saving against the agreed electricity-tariff benchmark."]),
        ("2. The Problem", "Common-area electricity cost keeps rising.", ["Lifts, pumps, lighting, clubhouse, STP, and other common loads add up every month.", "Many societies want solar but do not want to collect lakhs from residents."]),
        ("3. The Bright Roof Model", "We fund and maintain the system. The society gets lower power cost.", ["No upfront installation cost to the society.", "Existing grid supply remains for balance electricity.", "Society pays for solar electricity under the agreement."]),
        ("4. What The Society Provides", "First we verify fit. Then we propose.", ["12 months common-area electricity bills.", "Consumer service number and sanctioned load details.", "Rooftop and electrical-room access for survey."]),
        ("5. Process And Timeline", "Clear steps before any final commitment.", ["Bill and document review.", "Rooftop and electrical-room survey.", "Society-specific proposal.", "Committee/PPA review.", "Final approval and agreement."]),
        ("6. Today's Ask", "Permission for bill review and rooftop survey.", ["Review 12 months of bills.", "Inspect rooftop and electrical room by appointment.", "Prepare a site-specific proposal for committee review."]),
    ]
    for title, subtitle, bullets in slides:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph()
        p.add_run(subtitle).bold = True
        for b in bullets:
            add_bullet(doc, b)
    return save_doc(doc, "07-committee-pitch-deck-outline.docx")


def field_kit_person(name: str, role: str, filename: str) -> Path:
    doc = new_doc(f"{name} Field Kit", role)
    doc.add_heading("Primary Mission", level=1)
    if "Technical" in role:
        doc.add_paragraph("Make Bright Roof feel technically credible without turning the first meeting into an engineering session.")
        intro = "Hi [Name], this is Srinivas from Bright Roof Power Systems. We help apartment societies reduce common-area electricity bills through rooftop solar with no upfront installation cost to the society. My role is to help with rooftop suitability, electrical access, survey requirements, and installer coordination."
    elif "Relations" in role:
        doc.add_paragraph("Use local trust and apartment-network access to get serious committee conversations and document sharing.")
        intro = "Hi [Name], this is Vijay from Bright Roof Power Systems. We work with apartment societies in Hyderabad on rooftop solar where the society does not pay the installation cost upfront. The first step is only bill review and rooftop survey permission."
    else:
        doc.add_paragraph("Build and maintain a qualified pipeline of Hyderabad apartment-society leads.")
        intro = "Hi [Name], this is Kalyan from Bright Roof Power Systems. We help apartment societies reduce common-area electricity bills using rooftop solar, without the society paying upfront installation cost. Can I connect you with the Bright Roof team for a short discussion?"
    doc.add_heading("WhatsApp Intro", level=1)
    doc.add_paragraph(intro)
    doc.add_heading("Carry Pack", level=1)
    for item in ["Visiting card.", "One-page leaflet.", "Document request checklist.", "Consent form.", "Meeting notes template."]:
        add_bullet(doc, item)
    doc.add_heading("Escalate To Nrupal When", level=1)
    for item in ["A society asks for written proposal material.", "Bills are shared.", "A committee meeting is scheduled.", "Any number needs to go into writing.", "A stronger discount is requested."]:
        add_bullet(doc, item)
    return save_doc(doc, filename)


def all_docx() -> list[Path]:
    paths = [
        leaflet(),
        company_profile(),
        document_checklist(),
        consent_form(),
        faq(),
        comparison(),
        pitch_outline(),
        field_kit_person("Srinivas", "Technical & Project Coordination", "08-srinivas-field-kit.docx"),
        field_kit_person("Vijay", "Society Relations", "09-vijay-field-kit.docx"),
        field_kit_person("Kalyan", "Society Outreach", "10-kalyan-field-kit.docx"),
    ]
    return paths


def render_docx_to_pdf(docx_paths: list[Path]) -> None:
    if not RENDER.exists():
        print("render_docx.py not found; skipping PDF rendering", file=sys.stderr)
        return
    for path in docx_paths:
        out_dir = RENDER_DIR / path.stem
        out_dir.mkdir(parents=True, exist_ok=True)
        cmd = [str(PY), str(RENDER), str(path), "--output_dir", str(out_dir), "--emit_pdf", "--width", "1200", "--height", "1600"]
        subprocess.run(cmd, check=True)
        pdf = out_dir / f"{path.stem}.pdf"
        if pdf.exists():
            copyfile(pdf, PDF_DIR / pdf.name)


def excel_tracker() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Dashboard"
    leads = wb.create_sheet("Leads")
    review = wb.create_sheet("Weekly Review")
    lists = wb.create_sheet("Lists")

    stages = [
        "Lead identified", "Contact found", "Warm intro requested", "First contact made", "Interested",
        "Committee meeting planned", "Documents requested", "Bills received", "Survey permission requested",
        "Survey permission granted", "Proposal needed", "PPA review", "Won", "Lost", "Stalled",
    ]
    people = ["Srinivas", "Vijay", "Kalyan", "Nrupal"]

    headers = [
        "Lead ID", "Society name", "Area/locality", "Approx flat count", "Approx monthly common bill",
        "Contact person", "Contact role", "Phone", "WhatsApp available", "Warm intro source",
        "Assigned salesperson", "Stage", "Last contact date", "Next follow-up date", "Documents received",
        "Rooftop access status", "Main objection", "Next action", "Next action owner", "Notes",
    ]
    leads.append(headers)
    for col in range(1, len(headers) + 1):
        cell = leads.cell(1, col)
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    for i in range(2, 52):
        leads.cell(i, 1).value = f"BR-HYD-{i-1:03d}"
    widths = [14, 26, 18, 14, 20, 18, 18, 16, 14, 18, 18, 24, 16, 16, 22, 22, 26, 28, 18, 34]
    for idx, width in enumerate(widths, 1):
        leads.column_dimensions[chr(64 + idx) if idx <= 26 else "A"].width = width
    leads.freeze_panes = "A2"
    leads.auto_filter.ref = "A1:T51"

    for i, stage in enumerate(stages, 1):
        lists.cell(i, 1).value = stage
    for i, person in enumerate(people, 1):
        lists.cell(i, 2).value = person
    lists.sheet_state = "hidden"
    dv_stage = DataValidation(type="list", formula1="=Lists!$A$1:$A$15", allow_blank=True)
    dv_people = DataValidation(type="list", formula1="=Lists!$B$1:$B$4", allow_blank=True)
    leads.add_data_validation(dv_stage)
    leads.add_data_validation(dv_people)
    dv_stage.add("L2:L200")
    dv_people.add("K2:K200")
    dv_people.add("S2:S200")

    red_fill = PatternFill("solid", fgColor="FCE8E6")
    leads.conditional_formatting.add("N2:N200", FormulaRule(formula=['AND($N2<TODAY(),$L2<>"Won",$L2<>"Lost",$L2<>"Stalled",$N2<>"")'], fill=red_fill))

    ws["A1"] = "Bright Roof Sales Dashboard"
    ws["A1"].font = Font(size=18, bold=True, color=NAVY)
    ws["A2"] = "Use the Leads sheet as the working CRM. Dashboard formulas update from lead stages."
    ws["A4"] = "Metric"
    ws["B4"] = "Value"
    metrics = [
        ("Total leads", '=COUNTA(Leads!B2:B200)'),
        ("Interested", '=COUNTIF(Leads!L:L,"Interested")'),
        ("Committee meetings planned", '=COUNTIF(Leads!L:L,"Committee meeting planned")'),
        ("Bills received", '=COUNTIF(Leads!L:L,"Bills received")'),
        ("Survey permission granted", '=COUNTIF(Leads!L:L,"Survey permission granted")'),
        ("Won", '=COUNTIF(Leads!L:L,"Won")'),
        ("Stalled", '=COUNTIF(Leads!L:L,"Stalled")'),
    ]
    for r, (label, formula) in enumerate(metrics, 5):
        ws.cell(r, 1).value = label
        ws.cell(r, 2).value = formula
    ws.column_dimensions["A"].width = 34
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 4
    ws.column_dimensions["D"].width = 34
    ws.column_dimensions["E"].width = 10
    ws.column_dimensions["F"].width = 4
    for cell in ws[4]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(color=WHITE, bold=True)
    ws["D4"] = "Stage"
    ws["E4"] = "Count"
    for r, stage in enumerate(stages, 5):
        ws.cell(r, 4).value = stage
        ws.cell(r, 5).value = f'=COUNTIF(Leads!L:L,D{r})'
    for cell in ws[4][3:5]:
        cell.fill = PatternFill("solid", fgColor=AMBER)
        cell.font = Font(color=NAVY, bold=True)
    chart = BarChart()
    chart.title = "Lead Count By Stage"
    chart.height = 8
    chart.width = 14
    data = Reference(ws, min_col=5, min_row=4, max_row=19)
    cats = Reference(ws, min_col=4, min_row=5, max_row=19)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    ws.add_chart(chart, "G4")
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    review.append(["Week of", "", "Attendees", "Srinivas / Vijay / Kalyan / Nrupal"])
    review.append([])
    review.append(["Metric", "Target", "Actual"])
    for row in [
        ["New leads added", 8, ""],
        ["New contact persons found", 5, ""],
        ["First conversations", 3, ""],
        ["Committee meetings", 1, ""],
        ["Bill sets requested", 2, ""],
        ["Bills received", 1, ""],
        ["Survey permissions received", "0-1", ""],
    ]:
        review.append(row)
    review.append([])
    review.append(["Lead", "Owner", "Current stage", "Next action", "Due date"])
    for _ in range(12):
        review.append(["", "", "", "", ""])
    for sheet in [ws, leads, review]:
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = Border(bottom=Side(style="thin", color=LIGHT))
        sheet.sheet_view.showGridLines = False
    path = EXCEL_DIR / "bright-roof-lead-tracker.xlsx"
    wb.save(path)
    return path


def load_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
    ]
    for c in candidates:
        try:
            return ImageFont.truetype(c, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], font, fill, max_width: int, line_gap: int = 6) -> int:
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap
    return y


def visiting_cards() -> list[Path]:
    outputs = []
    for name, role, phone in PEOPLE:
        safe = name.lower().replace(" ", "-").replace("/", "-")
        front = Image.new("RGB", (1050, 600), f"#{CREAM}")
        draw = ImageDraw.Draw(front)
        navy = f"#{NAVY}"
        amber = f"#{AMBER}"
        slate = f"#{SLATE}"
        muted = f"#{MUTED}"
        draw.rectangle([0, 0, 1050, 86], fill=navy)
        draw.rectangle([70, 130, 176, 136], fill=amber)
        if LOGO.exists():
            logo = Image.open(LOGO).convert("RGBA")
            logo.thumbnail((180, 110))
            front.paste(logo, (70, 25), logo)
        draw.text((70, 168), name, font=load_font(50, True), fill=navy)
        draw.text((74, 230), role, font=load_font(28), fill=slate)
        draw.line([70, 310, 980, 310], fill=f"#{LIGHT}", width=3)
        y = 350
        for line in [phone, CONTACT["email"], CONTACT["website"]]:
            draw.text((70, y), line, font=load_font(26), fill=slate)
            y += 42
        draw.text((70, 530), "Bright Roof Power Systems  |  Solar that stays.", font=load_font(22, True), fill=navy)
        front_path = IMG_DIR / f"visiting-card-{safe}-front.png"
        front.save(front_path, dpi=(300, 300))
        outputs.append(front_path)

        back = Image.new("RGB", (1050, 600), navy)
        draw = ImageDraw.Draw(back)
        draw.rectangle([70, 75, 176, 81], fill=amber)
        draw.text((70, 116), "Rooftop solar for", font=load_font(42, True), fill=f"#{CREAM}")
        draw.text((70, 168), "apartment societies.", font=load_font(42, True), fill=f"#{CREAM}")
        y = 260
        y = draw_wrapped(draw, "No upfront installation cost to the society.", (70, y), load_font(30, True), amber, 850, 8)
        y = draw_wrapped(draw, "Bright Roof installs, owns, operates, and maintains the system.", (70, y + 22), load_font(26), f"#{CREAM}", 850, 7)
        y = draw_wrapped(draw, "Starting proposal: 10% saving against the agreed electricity benchmark.", (70, y + 18), load_font(26), f"#{CREAM}", 850, 7)
        draw.text((70, 528), f"{CONTACT['phone']}  |  {CONTACT['website']}", font=load_font(24, True), fill=amber)
        back_path = IMG_DIR / f"visiting-card-{safe}-back.png"
        back.save(back_path, dpi=(300, 300))
        outputs.append(back_path)
    return outputs


def main() -> None:
    ensure_dirs()
    docx_paths = all_docx()
    excel_tracker()
    visiting_cards()
    render_docx_to_pdf(docx_paths)
    print(f"Generated sales kit in {OUT}")


if __name__ == "__main__":
    main()
